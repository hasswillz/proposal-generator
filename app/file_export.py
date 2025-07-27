import os
from docx import Document
import pdfkit
from flask import current_app


class ProposalExporter:
    @staticmethod
    def export_pdf(content: str, filename: str, output_dir: str) -> str:
        """
        Exports the proposal content as a PDF using pdfkit
        Requires wkhtmltopdf to be installed on the system
        """
        try:
            filepath = os.path.join(output_dir, filename)

            # Configure pdfkit - update the path to your wkhtmltopdf installation
            config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')

            # Basic HTML template for the PDF
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 2cm; }}
                    h1 {{ color: #2c3e50; border-bottom: 1px solid #eee; }}
                    h2 {{ color: #34495e; }}
                    .header {{ margin-bottom: 2em; }}
                </style>
            </head>
            <body>
                {content}
            </body>
            </html>
            """

            options = {
                'encoding': 'UTF-8',
                'quiet': ''
            }

            pdfkit.from_string(html_content, filepath, configuration=config, options=options)
            return filepath
        except Exception as e:
            current_app.logger.error(f"PDF generation failed: {str(e)}")
            raise Exception("Failed to generate PDF. Please try again.")

    @staticmethod
    def export_docx(content: str, filename: str, output_dir: str) -> str:
        """
        Exports the proposal content as a DOCX file.
        Requires 'python-docx' library.
        """
        filepath = os.path.join(output_dir, filename)
        document = Document()

        # Add content to the document
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue

            if line.startswith('#'):
                # Handle headings
                level = line.count('#')
                heading_text = line.replace('#', '').strip()
                if level == 1:
                    document.add_heading(heading_text, level=1)
                elif level == 2:
                    document.add_heading(heading_text, level=2)
                else:
                    document.add_heading(heading_text, level=3)
            else:
                # Add regular paragraph
                paragraph = document.add_paragraph()
                paragraph.add_run(line)

        document.save(filepath)
        return filepath

    @staticmethod
    def export_markdown(content: str, filename: str, output_dir: str) -> str:
        """
        Exports the proposal content as a Markdown (.md) file.
        """
        filepath = os.path.join(output_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return filepath